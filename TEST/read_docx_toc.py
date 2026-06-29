import docx

doc = docx.Document("El-diputado-Nicolai-TOC.docx")
print("First 40 paragraphs of DOCX:")
for idx in range(min(40, len(doc.paragraphs))):
    p = doc.paragraphs[idx]
    print(f"[{idx}] (style: {p.style.name}): {repr(p.text)}")
